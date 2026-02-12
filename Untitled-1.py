import streamlit as st
import pandas as pd
from datetime import datetime
from streamlit_gsheets import GSheetsConnection
from docx import Document
from docx.shared import Pt
# Nota: docx2pdf geralmente não funciona em servidores Linux como o Streamlit Cloud. 
# Recomenda-se usar bibliotecas como 'fpdf' ou 'reportlab' para gerar PDFs diretamente.
import io
# No topo do seu código, onde você define a conexão
conn = st.connection("gsheets", type=GSheetsConnection)
# 1. Primeiro, defina os links das suas planilhas
# Substitua os links abaixo pelos links REAIS das suas planilhas
url_base_pessoas = "https://docs.google.com/spreadsheets/d/1AqdC3_qFiWvVkEunGBw9EuYRDiliMd9dORmZQNHZbVg/edit?gid=0#gid=0"
url_base_produtos = "https://docs.google.com/spreadsheets/d/1QwQYMfIflhy18xo2VGRYKnR9-uJJr69W_NfOln8KQJo/edit?gid=0#gid=0"
url_base_pedidos = "https://docs.google.com/spreadsheets/d/1U7FQYusJFAOoqRdp7bY3pODyCKGYnvDl9mLWNKAELoI/edit?usp=sharing"

# 2. Depois, crie a conexão
conn = st.connection("gsheets", type=GSheetsConnection)

# 3. Agora sim, faça a leitura (usando os nomes que você definiu acima)
dados_pessoas = conn.read(spreadsheet=url_base_pessoas, ttl=7200)
dados_produtos = conn.read(spreadsheet=url_base_produtos, ttl=60)
dados_pedidos = conn.read(spreadsheet=url_base_pedidos, ttl=20)
# --- CONFIGURAÇÃO DA CONEXÃO COM GOOGLE SHEETS ---
# Como você forneceu 3 URLs diferentes, criaremos uma função para ler cada uma
# --- FUNÇÕES AUXILIARES DE FORMATAÇÃO ---
def formatar_cpf_cnpj(doc):
    doc = ''.join(filter(str.isdigit, str(doc))) # Remove tudo que não é número
    if len(doc) == 11: # CPF
        return f"{doc[:3]}.{doc[3:6]}.{doc[6:9]}-{doc[9:]}"
    elif len(doc) == 14: # CNPJ
        return f"{doc[:2]}.{doc[2:5]}.{doc[5:8]}/{doc[8:12]}-{doc[12:]}"
    return doc

def formatar_cep(cep):
    cep = ''.join(filter(str.isdigit, str(cep)))
    if len(cep) == 8:
        return f"{cep[:5]}-{cep[5:]}"
    return cep

def formatar_telefone(tel):
    tel = ''.join(filter(str.isdigit, str(tel)))
    if len(tel) == 11: # Celular com DDD
        return f"({tel[:2]}) {tel[2:7]}-{tel[7:]}"
    elif len(tel) == 10: # Fixo com DDD
        return f"({tel[:2]}) {tel[2:6]}-{tel[6:]}"
    return tel
    
def carregar_dados_gsheets():
    conn = st.connection("gsheets", type=GSheetsConnection)
    
    try:
        # Carregando Base de Produtos
        url_produtos = "https://docs.google.com/spreadsheets/d/1QwQYMfIflhy18xo2VGRYKnR9-uJJr69W_NfOln8KQJo/edit?gid=0#gid=0"
        dados_produtos = conn.read(spreadsheet=url_produtos, ttl="5m")
        
        # Carregando Base de Pessoas
        url_pessoas = "https://docs.google.com/spreadsheets/d/1AqdC3_qFiWvVkEunGBw9EuYRDiliMd9dORmZQNHZbVg/edit?gid=0#gid=0"
        dados_pessoas = conn.read(spreadsheet=url_pessoas, ttl="5m")
        
        # Carregando Base de Pedidos
        url_pedidos = "https://docs.google.com/spreadsheets/d/1U7FQYusJFAOoqRdp7bY3pODyCKGYnvDl9mLWNKAELoI/edit?gid=0#gid=0"
        dados_pedidos = conn.read(spreadsheet=url_pedidos, ttl="0") # TTL 0 para ver pedidos novos na hora
        
        return dados_produtos, dados_pessoas, dados_pedidos, conn
    except Exception as e:
        st.error(f"Erro ao conectar com as planilhas: {e}")
        return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), None

# Inicializando as bases
dados_produtos, dados_pessoas, dados_pedidos, conn_gsheets = carregar_dados_gsheets()

# --- INTERFACE ---
st.sidebar.title("Menu")
pagina = st.sidebar.radio("Ir para:", [
    "Criar Pedido", 
    "Consultar Pedido", 
    "Consultar Produto", 
    "Cadastrar Produto", 
    "Cadastrar Pessoa", 
    "Consultar Pessoa", 
    "Formalizacao"
])

# Exemplo de como salvar um novo pedido (Para a Parte 7)
def salvar_novo_pedido(df_novo):
    url_pedidos = "https://docs.google.com/spreadsheets/d/1U7FQYusJFAOoqRdp7bY3pODyCKGYnvDl9mLWNKAELoI/edit?gid=0#gid=0"
    # Lê os existentes
    df_atual = conn_gsheets.read(spreadsheet=url_pedidos)
    # Concatena
    df_final = pd.concat([df_atual, df_novo], ignore_index=True)
    # Faz o update
    conn_gsheets.update(spreadsheet=url_pedidos, data=df_final)
    st.cache_data.clear()

if pagina == "Cadastrar Produto":
    st.title("Cadastro de Produtos")

    # URL da planilha de produtos para facilitar o uso no código
    url_base_produtos = "https://docs.google.com/spreadsheets/d/1QwQYMfIflhy18xo2VGRYKnR9-uJJr69W_NfOln8KQJo/edit?gid=0#gid=0"

    aba1, aba2 = st.tabs(["Cadastro Manual", "Importação em Massa (CSV)"])

    # --- ABA 1: CADASTRO MANUAL ---
    with aba1:
        st.info("O Valor Líquido será calculado: Custo + Impostos + Lucro")

        with st.form("form_cadastro"):
            st.subheader("1. Identificação")
            col_id_1, col_id_2 = st.columns(2)
            
            with col_id_1:
                id_sku = st.text_input("SKU / Código Interno (Obrigatório)")
                descricao = st.text_input("Descrição do Produto (Obrigatório)")
                marca = st.text_input("Marca / Fabricante")
                
            with col_id_2:
                categoria = st.selectbox("Categoria", ["Geral", "Eletrônicos", "Vestuário", "Ferramentas", "Outros"])
                fornecedor = st.selectbox("Fornecedor", ["Samsung", "Apple", "LG", "Motorola", "Outros"])
                
                c_est1, c_est2 = st.columns(2)
                estoque_atual = c_est1.number_input("Estoque Atual", min_value=0, step=1)
                estoque_minimo = c_est2.number_input("Estoque Mínimo", min_value=1, value=5)

            st.divider() 

            st.subheader("2. Custos e Precificação")
            col_fin_1, col_fin_2, col_fin_3 = st.columns(3)
            with col_fin_1:
                preco_custo = st.number_input("Preço de Custo (R$)", min_value=0.00, step=0.01)
                lucro = st.number_input("Margem de Lucro (R$)", min_value=0.00, step=0.01)
            
            with col_fin_2:
                icms = st.number_input("ICMS (R$)", min_value=0.0, step=0.01)
                ipi = st.number_input("IPI (R$)", min_value=0.0, step=0.01)
            
            with col_fin_3:
                valor_st = st.number_input("ST (R$)", min_value=0.0, step=0.01)
                ncm = st.text_input("NCM")

            botao_salvar = st.form_submit_button("Salvar Produto")

        # --- LÓGICA DE SALVAR (MANUAL) ---
        if botao_salvar:
            erros = []
            if not id_sku: erros.append("O SKU é obrigatório.")
            if not descricao: erros.append("A Descrição é obrigatória.")
            
            # Validação de Duplicados no Google Sheets
            if not dados_produtos.empty:
                lista_skus = dados_produtos["id_sku"].astype(str).tolist()
                if str(id_sku) in lista_skus:
                    erros.append(f"ERRO CRÍTICO: O SKU '{id_sku}' já existe no sistema!")

            if len(erros) > 0:
                for erro in erros: st.error(erro)
            else:
                valor_liquido = preco_custo + icms + ipi + valor_st + lucro
                
                nova_linha = pd.DataFrame({
                    "id_sku": [id_sku],
                    "descricao": [descricao],
                    "categoria": [categoria],
                    "marca": [marca],
                    "fornecedor": [fornecedor],
                    "ncm": [ncm],
                    "preco_custo": [preco_custo],
                    "lucro": [lucro],
                    "icms": [icms], "ipi": [ipi], "st": [valor_st],
                    "valor_liquido": [valor_liquido],
                    "estoque_atual": [estoque_atual],
                    "estoque_minimo": [estoque_minimo],
                    "ativo": [True],
                    "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
                })

                # Gravação no Google Sheets
                try:
                    df_atualizado = pd.concat([dados_produtos, nova_linha], ignore_index=True)
                    conn_gsheets.update(spreadsheet=url_base_produtos, data=df_atualizado)
                    st.success(f"✅ Produto {id_sku} cadastrado no Google Sheets!")
                    st.cache_data.clear() # Limpa o cache para atualizar a consulta
                    st.rerun()
                except Exception as e:
                    st.error(f"Erro ao salvar na planilha: {e}")

    # --- ABA 2: IMPORTAÇÃO EM MASSA (CSV) ---
    with aba2:
        st.header("Importar Produtos via CSV")
        
        with st.expander("📖 Instruções"):
            st.markdown("O arquivo deve ser **CSV (ponto e vírgula)** com as colunas idênticas à planilha.")
        
        arquivo_upload = st.file_uploader("Arraste seu arquivo CSV aqui", type=["csv"])
        
        if arquivo_upload is not None:
            try:
                df_novo = pd.read_csv(arquivo_upload, sep=";")
                st.dataframe(df_novo.head())
                
                if st.button("Confirmar Importação em Massa"):
                    skus_existentes = dados_produtos["id_sku"].astype(str).tolist()
                    df_novo_filtrado = df_novo[~df_novo["id_sku"].astype(str).isin(skus_existentes)]
                    
                    if not df_novo_filtrado.empty:
                        if "data_cadastro" not in df_novo_filtrado.columns:
                            df_novo_filtrado["data_cadastro"] = datetime.now().strftime("%d/%m/%Y")
                        
                        # Atualização no Google Sheets
                        dados_atualizados = pd.concat([dados_produtos, df_novo_filtrado], ignore_index=True)
                        conn_gsheets.update(spreadsheet=url_base_produtos, data=dados_atualizados)
                        
                        st.success(f"✅ {len(df_novo_filtrado)} produtos importados com sucesso!")
                        st.cache_data.clear()
                        st.rerun()
                    else:
                        st.warning("Todos os produtos do arquivo já existem na base.")
            except Exception as e:
                st.error(f"Erro no processamento: {e}")
elif pagina == "Consultar Produto":
    st.title("Consulta de Produtos")

    # Criamos duas colunas: uma estreita para os filtros e uma larga para o resultado
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros de Busca")
        # Busca por SKU
        filtro_sku = st.text_input("Código SKU")
        
        # Busca por Descrição
        filtro_desc = st.text_input("Descrição do Produto")
        
        st.divider()
        if st.button("🔄 Atualizar Dados"):
            st.cache_data.clear()
            st.rerun()
        st.caption("Dica: A busca encontra palavras parciais e ignora maiúsculas/minúsculas.")

    with col_resultado:
        # Usamos os dados carregados na Parte 1 via conn_gsheets
        if dados_produtos.empty:
            st.warning("A base de produtos está vazia ou não foi carregada corretamente.")
        else:
            df_filtrado = dados_produtos.copy()

            # Lógica de Filtro em Tempo Real
            if filtro_sku:
                df_filtrado = df_filtrado[df_filtrado["id_sku"].astype(str).str.contains(filtro_sku, case=False, na=False)]
            
            if filtro_desc:
                df_filtrado = df_filtrado[df_filtrado["descricao"].astype(str).str.contains(filtro_desc, case=False, na=False)]

            # Colunas para exibição (Exatamente como estão no Google Sheets)
            colunas_exibicao = [
                "id_sku", 
                "descricao", 
                "fornecedor", 
                "preco_custo", 
                "lucro", 
                "valor_liquido"
            ]
            
            try:
                # Filtrar colunas existentes
                exibicao = df_filtrado[colunas_exibicao].copy()
                
                # Garantir que valores financeiros sejam numéricos para exibição correta
                cols_financeiras = ["preco_custo", "lucro", "valor_liquido"]
                for col in cols_financeiras:
                    exibicao[col] = pd.to_numeric(exibicao[col], errors='coerce').fillna(0.0)
                
                # Renomear colunas para a tabela amigável
                exibicao.columns = ["SKU", "DESCRIÇÃO", "FORNECEDOR", "CUSTO (R$)", "LUCRO (R$)", "VALOR LÍQUIDO (R$)"]
                
                st.subheader(f"Resultados ({len(exibicao)} encontrados)")
                
                if len(exibicao) > 0:
                    st.dataframe(
                        exibicao.style.format({
                            "CUSTO (R$)": "{:.2f}",
                            "LUCRO (R$)": "{:.2f}",
                            "VALOR LÍQUIDO (R$)": "{:.2f}"
                        }), 
                        use_container_width=True, 
                        hide_index=True
                    )
                else:
                    st.info("Nenhum produto encontrado com os filtros aplicados.")
                    
            except KeyError as e:
                st.error(f"Erro: Coluna não encontrada na planilha: {e}")
                st.write("Colunas disponíveis na sua planilha:", list(dados_produtos.columns))

# --- INÍCIO DA PÁGINA ---
elif pagina == "Cadastrar Pessoa":
    st.title("Cadastro de Clientes e Fornecedores")
    url_base_pessoas = "https://docs.google.com/spreadsheets/d/1AqdC3_qFiWvVkEunGBw9EuYRDiliMd9dORmZQNHZbVg/edit?gid=0#gid=0"

    if dados_pessoas.empty:
        dados_pessoas = pd.DataFrame(columns=[
            "id_documento", "tipo_pessoa", "nome_razao", "nome_fantasia", 
            "rg_ie", "email", "telefone", "cep", "endereco", "numero", 
            "complemento", "bairro", "cidade", "uf", "categoria", 
            "limite_credito", "status", "data_cadastro"
        ])

    with st.form("form_pessoas", clear_on_submit=False):
        st.subheader("1. Identificação Principal")
        col_id_1, col_id_2, col_id_3 = st.columns([2, 2, 2])
        
        with col_id_1:
            tipo_pessoa = st.selectbox("Tipo de Pessoa", ["Física", "Jurídica"])
            label_doc = "CPF (Somente números)" if tipo_pessoa == "Física" else "CNPJ (Somente números)"
            id_documento_raw = st.text_input(label_doc, help="A formatação será aplicada automaticamente ao salvar.")
            
        with col_id_2:
            categoria = st.selectbox("Categoria", ["Cliente", "Fornecedor", "Transportadora", "Ambos"])
            status = st.selectbox("Status Inicial", ["Ativo", "Inativo", "Bloqueado"])
            
        with col_id_3:
            limite_credito = st.number_input("Limite de Crédito (R$)", min_value=0.0, step=100.0)

        st.divider()
        st.subheader("2. Dados Pessoais / Empresariais")
        col_dados_1, col_dados_2 = st.columns(2)
        
        with col_dados_1:
            label_nome = "Nome Completo" if tipo_pessoa == "Física" else "Razão Social"
            nome_razao = st.text_input(label_nome)
            nome_fantasia = st.text_input("Nome Fantasia (Se houver)")
            
        with col_dados_2:
            label_rg = "RG" if tipo_pessoa == "Física" else "Inscrição Estadual"
            rg_ie = st.text_input(label_rg)
            email = st.text_input("E-mail para contato/NFe")
            telefone_raw = st.text_input("WhatsApp / Telefone", help="Ex: 11999998888")

        st.divider()
        st.subheader("3. Endereço")
        col_end_1, col_end_2, col_end_3 = st.columns([1, 2, 1])
        with col_end_1:
            cep_raw = st.text_input("CEP", help="Ex: 01234567")
            numero = st.text_input("Número")
        with col_end_2:
            endereco = st.text_input("Logradouro (Rua/Av)")
            complemento = st.text_input("Complemento")
        with col_end_3:
            bairro = st.text_input("Bairro")
            cidade = st.text_input("Cidade")
            uf = st.selectbox("UF", ["AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR", "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO"])

        botao_salvar_pessoa = st.form_submit_button("Finalizar Cadastro", use_container_width=True)

    # --- LÓGICA DE SALVAR E FORMATAR ---
    if botao_salvar_pessoa:
        # Aplicando as formatações antes de validar
        id_documento = formatar_cpf_cnpj(id_documento_raw)
        telefone = formatar_telefone(telefone_raw)
        cep = formatar_cep(cep_raw)
        
        erros_pessoa = []
        
        # Validação de comprimento mínimo (para evitar salvar lixo)
        if tipo_pessoa == "Física" and len(''.join(filter(str.isdigit, id_documento_raw))) != 11:
            erros_pessoa.append("CPF inválido. Deve conter 11 dígitos.")
        if tipo_pessoa == "Jurídica" and len(''.join(filter(str.isdigit, id_documento_raw))) != 14:
            erros_pessoa.append("CNPJ inválido. Deve conter 14 dígitos.")
        if not nome_razao: erros_pessoa.append(f"O campo {label_nome} é obrigatório.")
        
        # Trava de Duplicidade
        if not dados_pessoas.empty:
            if id_documento in dados_pessoas["id_documento"].astype(str).tolist():
                erros_pessoa.append(f"O documento {id_documento} já existe na base!")

        if erros_pessoa:
            for erro in erros_pessoa:
                st.error(erro)
        else:
            nova_pessoa = pd.DataFrame({
                "id_documento": [id_documento],
                "tipo_pessoa": [tipo_pessoa],
                "nome_razao": [nome_razao],
                "nome_fantasia": [nome_fantasia],
                "rg_ie": [rg_ie],
                "email": [email],
                "telefone": [telefone],
                "cep": [cep],
                "endereco": [endereco],
                "numero": [numero],
                "complemento": [complemento],
                "bairro": [bairro],
                "cidade": [cidade],
                "uf": [uf],
                "categoria": [categoria],
                "limite_credito": [limite_credito],
                "status": [status],
                "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
            })

            try:
                dados_atualizados = pd.concat([dados_pessoas, nova_pessoa], ignore_index=True)
                conn_gsheets.update(spreadsheet=url_base_pessoas, data=dados_atualizados)
                
                st.success(f"✅ {nome_razao} cadastrado(a) com sucesso!")
                st.cache_data.clear()
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao salvar na planilha: {e}")
elif pagina == "Consultar Pessoa":
    st.title("Consulta de Clientes / Fornecedores")

    # 1. Verificar se os dados foram carregados (dados_pessoas já vem da Parte 1)
    if dados_pessoas.empty:
        st.warning("Nenhuma base de pessoas encontrada. Cadastre alguém primeiro ou verifique a conexão!")
        if st.button("🔄 Tentar Recarregar"):
            st.cache_data.clear()
            st.rerun()
        st.stop() 

    # 2. Layout de Colunas
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros")
        filtro_doc = st.text_input("Buscar por CPF/CNPJ")
        filtro_nome = st.text_input("Buscar por Nome/Razão")
        
        filtro_cat = st.multiselect(
            "Filtrar Categoria", 
            ["Cliente", "Fornecedor", "Transportadora", "Ambos"],
            default=[]
        )
        
        st.divider()
        if st.button("🔄 Atualizar Planilha"):
            st.cache_data.clear()
            st.rerun()
        st.caption("A busca por nome ignora maiúsculas e minúsculas.")

    with col_resultado:
        # Criamos a cópia para filtrar
        df_p_filtrado = dados_pessoas.copy()

        # Lógica de Filtro em Tempo Real
        if filtro_doc:
            # Garantimos que o documento seja tratado como string para a busca parcial
            df_p_filtrado = df_p_filtrado[df_p_filtrado["id_documento"].astype(str).str.contains(filtro_doc, na=False)]
        
        if filtro_nome:
            df_p_filtrado = df_p_filtrado[df_p_filtrado["nome_razao"].astype(str).str.contains(filtro_nome, case=False, na=False)]
        
        if filtro_cat:
            # Filtra se a categoria está na lista selecionada no multiselect
            df_p_filtrado = df_p_filtrado[df_p_filtrado["categoria"].isin(filtro_cat)]

        # Seleção de Colunas para a Tabela
        colunas_ver = [
            "id_documento",
            "nome_razao",
            "categoria",
            "email",
            "telefone",
            "cidade",
            "status"
        ]

        try:
            # Selecionamos apenas as colunas desejadas para exibição
            exibicao_p = df_p_filtrado[colunas_ver].copy()
            
            # Renomeando para ficar apresentável
            exibicao_p.columns = ["DOCUMENTO", "NOME / RAZÃO SOCIAL", "CATEGORIA", "E-MAIL", "CONTATO", "CIDADE", "STATUS"]

            st.subheader(f"Registros Encontrados ({len(exibicao_p)})")
            
            if len(exibicao_p) > 0:
                st.dataframe(
                    exibicao_p, 
                    use_container_width=True, 
                    hide_index=True
                )
                
                # Feedback visual caso haja apenas um resultado
                if len(exibicao_p) == 1:
                    st.success("💡 Registro localizado com sucesso!")
            else:
                st.info("Nenhuma pessoa encontrada com esses critérios.")
                
        except KeyError as e:
            st.error(f"Erro: Alguma coluna não existe na planilha 'Base_Pessoas': {e}")
            st.write("Colunas encontradas:", list(dados_pessoas.columns))
elif pagina == "Criar Pedido":
    st.title("Central de Pedidos")
    url_base_pedidos = "https://docs.google.com/spreadsheets/d/1U7FQYusJFAOoqRdp7bY3pODyCKGYnvDl9mLWNKAELoI/edit?gid=0#gid=0"

    # --- INICIALIZAÇÃO DE ESTADOS ---
    if "carrinho" not in st.session_state: st.session_state.carrinho = []
    if "cliente_selecionado" not in st.session_state: st.session_state.cliente_selecionado = None
    if "produto_selecionado" not in st.session_state: st.session_state.produto_selecionado = None

    # Lógica de ID Sequencial usando os dados carregados na Parte 1
    if not dados_pedidos.empty:
        # Garantir que o ID seja numérico para achar o máximo
        ids = pd.to_numeric(dados_pedidos["id_pedido"], errors='coerce').fillna(0)
        proximo_id = int(ids.max()) + 1
    else:
        proximo_id = 1

    st.subheader(f"Pedido Nº: {proximo_id}")

    # --- FUNÇÕES DE BUSCA (DIALOGS) ---
    @st.dialog("Buscar Cliente")
    def buscar_cliente_pop():
        st.write("Pesquise e selecione o cliente.")
        filtro = st.text_input("Nome ou CPF/CNPJ")
        if not dados_pessoas.empty:
            df_p = dados_pessoas.copy()
            if filtro:
                df_p = df_p[df_p["nome_razao"].str.contains(filtro, case=False, na=False) | 
                            df_p["id_documento"].astype(str).str.contains(filtro, na=False)]
            
            for _, row in df_p.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['nome_razao']}** ({row['id_documento']})")
                if col2.button("Selecionar", key=f"sel_p_{row['id_documento']}"):
                    st.session_state.cliente_selecionado = row.to_dict()
                    st.rerun()
        else:
            st.error("Base de Pessoas vazia ou não carregada!")

    @st.dialog("Buscar Produto")
    def buscar_produto_pop():
        st.write("Pesquise o SKU ou Descrição")
        filtro = st.text_input("Palavra-chave")
        if not dados_produtos.empty:
            df_prod = dados_produtos.copy()
            if filtro:
                df_prod = df_prod[df_prod["descricao"].str.contains(filtro, case=False, na=False) | 
                                  df_prod["id_sku"].astype(str).str.contains(filtro, na=False)]
            
            for _, row in df_prod.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['id_sku']}** - {row['descricao']}")
                if col2.button("Selecionar", key=f"sel_prod_{row['id_sku']}"):
                    st.session_state.produto_selecionado = row.to_dict()
                    st.rerun()

    # --- ÁREA 1: IDENTIFICAÇÃO DO CLIENTE ---
    with st.container(border=True):
        col_cli_1, col_cli_2 = st.columns([3, 1])
        with col_cli_1:
            doc_exibicao = st.session_state.cliente_selecionado['id_documento'] if st.session_state.cliente_selecionado else ""
            st.text_input("Cliente selecionado:", value=doc_exibicao, disabled=True)
        with col_cli_2:
            st.write("##")
            if st.button("🔍 Buscar Cliente", use_container_width=True):
                buscar_cliente_pop()

        if st.session_state.cliente_selecionado:
            c = st.session_state.cliente_selecionado
            st.success(f"**{c['nome_razao']}** | Limite: R$ {c.get('limite_credito', 0)}")

    # --- ÁREA 2: INCLUSÃO DE PRODUTOS ---
    with st.container(border=True):
        st.write("### Adicionar Itens")
        col_prod_1, col_prod_2, col_prod_3 = st.columns([2, 1, 1])
        
        with col_prod_1:
            sku_exibicao = st.session_state.produto_selecionado['id_sku'] if st.session_state.produto_selecionado else ""
            st.text_input("SKU selecionado:", value=sku_exibicao, disabled=True)
        with col_prod_2:
            st.write("##")
            if st.button("🔍 Buscar SKU", use_container_width=True):
                buscar_produto_pop()
        with col_prod_3:
            qtd = st.number_input("Quantidade", min_value=1, value=1)

        if st.session_state.produto_selecionado:
            p = st.session_state.produto_selecionado
            v_liq = pd.to_numeric(p['valor_liquido'], errors='coerce') or 0.0
            
            col_desc_1, col_desc_2 = st.columns(2)
            desconto = col_desc_1.number_input("Desconto (R$) / Negativo p/ Acréscimo", value=0.0)
            
            valor_final_item = v_liq - desconto
            col_desc_2.metric("Preço Unit. Final", f"R$ {valor_final_item:.2f}")

            if st.button("➕ Adicionar ao Carrinho", use_container_width=True):
                st.session_state.carrinho.append({
                    "sku": p['id_sku'],
                    "descricao": p['descricao'],
                    "qtd": qtd,
                    "valor_unit": valor_final_item,
                    "subtotal": valor_final_item * qtd
                })
                st.session_state.produto_selecionado = None
                st.rerun()

    # --- ÁREA 3: REVISÃO E FINALIZAÇÃO ---
    if st.session_state.carrinho:
        st.divider()
        st.subheader("🛒 Resumo do Carrinho")
        df_carrinho = pd.DataFrame(st.session_state.carrinho)
        st.table(df_carrinho[["sku", "descricao", "qtd", "valor_unit", "subtotal"]])
        
        if st.button("🗑️ Esvaziar Carrinho"):
            st.session_state.carrinho = []
            st.rerun()

        subtotal_itens = df_carrinho["subtotal"].sum()

        # FORMULÁRIO FINAL: Para evitar o bug do Enter salvando sozinho,
        # validamos o clique do botão explicitamente.
        with st.form("form_finalizar"):
            st.write("### Finalização")
            f1, f2, f3 = st.columns([2, 1, 1])
            
            tipo = f1.selectbox("Tipo de Documento", ["ORÇAMENTO", "PEDIDO", "COTAÇÃO"])
            frete = f2.number_input("Frete Total (R$)", min_value=0.0, step=5.0)
            obs = st.text_area("Observações do Pedido")
            
            total_geral = subtotal_itens + frete
            f3.metric("TOTAL GERAL", f"R$ {total_geral:.2f}")
            
            # Trava de segurança no botão
            confirmar = st.form_submit_button("💾 CONFIRMAR E SALVAR NO GOOGLE SHEETS", use_container_width=True)
            
            if confirmar:
                if not st.session_state.cliente_selecionado:
                    st.error("❌ Erro: Você esqueceu de selecionar o cliente!")
                else:
                    novas_linhas = []
                    for item in st.session_state.carrinho:
                        novas_linhas.append({
                            "id_pedido": proximo_id,
                            "data_pedido": datetime.now().strftime("%d/%m/%Y %H:%M"),
                            "doc_cliente": str(st.session_state.cliente_selecionado['id_documento']),
                            "nome_cliente": st.session_state.cliente_selecionado['nome_razao'],
                            "sku_item": item['sku'],
                            "qtd": item['qtd'],
                            "valor_final": item['valor_unit'],
                            "frete": frete,
                            "tipo": tipo,
                            "observacao": obs
                        })
                    
                    try:
                        df_novo_pedido = pd.DataFrame(novas_linhas)
                        # Combina com os dados que já existem na planilha
                        df_completo = pd.concat([dados_pedidos, df_novo_pedido], ignore_index=True)
                        conn_gsheets.update(spreadsheet=url_base_pedidos, data=df_completo)
                        
                        st.balloons()
                        st.success(f"✅ Pedido Nº {proximo_id} salvo com sucesso!")
                        st.session_state.carrinho = []
                        st.session_state.cliente_selecionado = None
                        st.cache_data.clear()
                        # Pequena pausa antes do rerun para o usuário ver a mensagem
                        import time
                        time.sleep(2)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erro ao salvar: {e}")
elif pagina == "Consultar Pedido":
    st.title("Gestão e Consulta de Pedidos")

    # 1. Padronização Inicial dos Dados
    df_pedidos = dados_pedidos.copy()
    df_produtos = dados_produtos.copy()
    df_pessoas = dados_pessoas.copy()

    # Garantir tipos corretos para evitar erro de merge
    df_pedidos['id_pedido'] = pd.to_numeric(df_pedidos['id_pedido'], errors='coerce').fillna(0).astype(int)
    df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
    df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)

    # Identificar nome da coluna de frete (evita KeyError)
    col_frete = 'frete_total' if 'frete_total' in df_pedidos.columns else 'frete'

    # --- FUNÇÃO DO POP-UP DE EDIÇÃO ---
    @st.dialog("Editar Pedido", width="large")
    def editar_pedido_pop(id_p):
        st.write(f"### Ajustes no Pedido #{id_p}")
        
        # Filtra itens atuais do pedido
        mask = df_pedidos["id_pedido"] == id_p
        itens_atuais = df_pedidos[mask].copy()
        
        # Inicializa o estado de edição se não existir
        if "edit_carrinho" not in st.session_state:
            st.session_state.edit_carrinho = itens_atuais.to_dict('records')
            st.session_state.edit_tipo = str(itens_atuais.iloc[0]['tipo'])
            st.session_state.edit_obs = str(itens_atuais.iloc[0]['observacao'])
            st.session_state.edit_frete = float(itens_atuais.iloc[0][col_frete])

        # --- SEÇÃO A: CABEÇALHO DO PEDIDO ---
        c1, c2, c3 = st.columns([2, 1, 1])
        with c1:
            novo_tipo = st.selectbox("Tipo", ["ORÇAMENTO", "PEDIDO", "COTAÇÃO"], 
                                    index=["ORÇAMENTO", "PEDIDO", "COTAÇÃO"].index(st.session_state.edit_tipo),
                                    key="edit_tipo_sel")
        with c2:
            novo_frete = st.number_input("Frete Total", value=st.session_state.edit_frete, key="edit_frete_val")
        with c3:
            st.write("##")
            if st.button("Limpar Carrinho", type="secondary"):
                st.session_state.edit_carrinho = []
                st.rerun()

        nova_obs = st.text_area("Observações", value=st.session_state.edit_obs, key="edit_obs_val")

        st.divider()

        # --- SEÇÃO B: ADICIONAR NOVOS ITENS ---
        st.subheader("Adicionar Novo Item")
        col_add_1, col_add_2, col_add_3 = st.columns([3, 1, 1])
        
        with col_add_1:
            # Lista de produtos da Base de Dados
            lista_prod = df_produtos['descricao'].tolist()
            prod_sel = st.selectbox("Produto", lista_prod, key="sel_prod_edit")
        with col_add_2:
            qtd_add = st.number_input("Qtd", min_value=1, value=1, key="qtd_prod_edit")
        with col_add_3:
            st.write("##")
            if st.button("➕ Add", use_container_width=True):
                # Busca info do produto
                info_p = df_produtos[df_produtos['descricao'] == prod_sel].iloc[0]
                novo_item = {
                    "id_pedido": id_p,
                    "data_pedido": itens_atuais.iloc[0]['data_pedido'],
                    "doc_cliente": itens_atuais.iloc[0]['doc_cliente'],
                    "nome_cliente": itens_atuais.iloc[0]['nome_cliente'],
                    "sku_item": str(info_p['id_sku']),
                    "qtd": qtd_add,
                    "valor_final": float(info_p['valor_liquido']),
                    col_frete: novo_frete,
                    "tipo": novo_tipo,
                    "observacao": nova_obs
                }
                st.session_state.edit_carrinho.append(novo_item)
                st.rerun()

        st.divider()

        # --- SEÇÃO C: LISTA DE ITENS ATUAIS ---
        st.subheader("Itens no Pedido")
        if not st.session_state.edit_carrinho:
            st.warning("O pedido está vazio!")
        else:
            for i, item in enumerate(st.session_state.edit_carrinho):
                # Busca descrição para exibir (merge manual)
                desc = df_produtos[df_produtos['id_sku'] == item['sku_item']]['descricao'].values
                nome_p = desc[0] if len(desc) > 0 else "Produto não encontrado"
                
                cc1, cc2, cc3, cc4 = st.columns([3, 1, 1, 0.5])
                cc1.caption(f"**{item['sku_item']}** - {nome_p}")
                cc2.write(f"Qtd: {item['qtd']}")
                cc3.write(f"R$ {float(item['valor_final']):.2f}")
                # Botão de remover com chave única (DuplicateKey Fix)
                if cc4.button("🗑️", key=f"del_it_{i}_{item['sku_item']}"):
                    st.session_state.edit_carrinho.pop(i)
                    st.rerun()

        st.divider()

        # --- SEÇÃO D: SALVAR NO BANCO ---
        if st.button("💾 SALVAR ALTERAÇÕES", type="primary", use_container_width=True):
            if not st.session_state.edit_carrinho:
                st.error("Não é possível salvar um pedido sem itens.")
            else:
                # 1. Remove itens antigos do ID atual
                df_limpo = df_pedidos[df_pedidos["id_pedido"] != id_p]
                
                # 2. Prepara os novos dados (garantindo que cabeçalhos batam)
                novos_itens_df = pd.DataFrame(st.session_state.edit_carrinho)
                
                # Atualiza os campos globais nos itens (tipo, frete e obs podem ter mudado)
                novos_itens_df[col_frete] = novo_frete
                novos_itens_df['tipo'] = novo_tipo
                novos_itens_df['observacao'] = nova_obs
                
                # 3. Une tudo e salva
                df_final = pd.concat([df_limpo, novos_itens_df], ignore_index=True)
                
                try:
                    # Se estiver usando CSV Local:
                    df_final.to_csv("Base_Pedido.csv", sep=";", index=False)
                    # Se estiver usando Google Sheets (Descomente abaixo):
                    # conn_gsheets.update(spreadsheet=url_base_pedidos, data=df_final)
                    
                    st.success("Pedido Atualizado com Sucesso!")
                    del st.session_state.edit_carrinho # Limpa estado
                    st.cache_data.clear()
                    st.rerun()
                except Exception as e:
                    st.error(f"Erro ao salvar: {e}")

    # --- LAYOUT PRINCIPAL DA PÁGINA ---
    col_filtros, col_detalhe = st.columns([1, 2.5])

    with col_filtros:
        st.subheader("Filtros")
        f_id_check = st.checkbox("Nº do Pedido")
        f_id = st.number_input("ID", min_value=1, step=1, disabled=not f_id_check)
        
        f_cliente_check = st.checkbox("CPF/CNPJ")
        f_cliente = st.text_input("Documento", disabled=not f_cliente_check)

        df_f = df_pedidos.copy()
        if f_id_check: df_f = df_f[df_f["id_pedido"] == f_id]
        if f_cliente_check: df_f = df_f[df_f["doc_cliente"].str.contains(f_cliente, na=False)]

        st.divider()
        lista_ids = sorted(df_f["id_pedido"].unique(), reverse=True)
        if lista_ids:
            id_selecionado = st.selectbox("Selecione o Pedido", lista_ids)
        else:
            st.warning("Nenhum pedido encontrado.")
            id_selecionado = None

    with col_detalhe:
        if id_selecionado:
            # Filtra e exibe resumo
            itens_venda = df_pedidos[df_pedidos["id_pedido"] == id_selecionado]
            
            c_header, c_edit = st.columns([3, 1])
            c_header.subheader(f"Pedido #{id_selecionado}")
            
            # Botão para abrir o Pop-up
            if c_edit.button("📝 Editar Pedido", use_container_width=True):
                if "edit_carrinho" in st.session_state: del st.session_state.edit_carrinho
                editar_pedido_pop(id_selecionado)

            with st.container(border=True):
                st.markdown(f"**Cliente:** {itens_venda.iloc[0]['nome_cliente']}")
                st.markdown(f"**Tipo:** `{itens_venda.iloc[0]['tipo']}` | **Data:** {itens_venda.iloc[0]['data_pedido']}")
                
                # Tabela de Itens
                st.dataframe(itens_venda[["sku_item", "qtd", "valor_final"]], use_container_width=True, hide_index=True)
                
                # Cálculo do Total
                valor_itens = (itens_venda["qtd"].astype(float) * itens_venda["valor_final"].astype(float)).sum()
                v_frete = float(itens_venda.iloc[0][col_frete])
                
                col_m1, col_m2 = st.columns(2)
                col_m1.metric("Subtotal Itens", f"R$ {valor_itens:.2f}")
                col_m2.metric("Total Geral (c/ Frete)", f"R$ {valor_itens + v_frete:.2f}")
                
                if str(itens_venda.iloc[0]['observacao']).lower() != 'nan':
                    st.info(f"**Obs:** {itens_venda.iloc[0]['observacao']}")
        else:
            st.info("Utilize os filtros à esquerda para localizar um pedido.")
elif pagina == "Formalizacao":
    st.title("📄 Formalização de Proposta")
    
    # 1. Uso dos dados globais carregados na Parte 1
    if dados_pedidos.empty:
        st.warning("Nenhum pedido encontrado na base para formalizar.")
        st.stop()

    df_pedidos = dados_pedidos.copy()
    df_produtos = dados_produtos.copy()
    df_pessoas = dados_pessoas.copy()

    # Padronização de tipos para o merge não falhar
    df_pedidos['id_pedido'] = pd.to_numeric(df_pedidos['id_pedido'], errors='coerce').fillna(0).astype(int)
    df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
    df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
    df_pessoas['id_documento'] = df_pessoas['id_documento'].astype(str)
    
    lista_pedidos = sorted(df_pedidos["id_pedido"].unique(), reverse=True)

    # 2. Seleção do Pedido
    id_escolhido = st.selectbox("Selecione o Número do Pedido", lista_pedidos, index=None, placeholder="Escolha um pedido para gerar o documento...")

    if id_escolhido:
        # Filtra dados do Pedido selecionado
        dados_venda = df_pedidos[df_pedidos["id_pedido"] == id_escolhido]
        doc_cliente = str(dados_venda.iloc[0]["doc_cliente"])
        
        # Merge com produtos para pegar descrição e marca
        itens_completos = dados_venda.merge(
            df_produtos[['id_sku', 'descricao', 'marca', 'preco_custo']], 
            left_on='sku_item', right_on='id_sku', how='left'
        )
        
        # Busca dados do Cliente na base de Pessoas
        filtro_cliente = df_pessoas[df_pessoas["id_documento"] == doc_cliente]
        
        if filtro_cliente.empty:
            st.error("Cliente deste pedido não encontrado na Base de Pessoas.")
            st.stop()
            
        cliente_info = filtro_cliente.iloc[0]

        # --- PAINEL DE CONFERÊNCIA ANTES DE GERAR ---
        with st.container(border=True):
            st.subheader(f"Resumo para Documento: Pedido #{id_escolhido}")
            c1, c2 = st.columns(2)
            with c1:
                st.write(f"**Cliente:** {cliente_info['nome_razao']}")
                st.write(f"**Data Original:** {dados_venda.iloc[0]['data_pedido']}")
                st.caption(f"📍 Entrega: {cliente_info['cidade']}/{cliente_info['uf']}")
            with c2:
                # Cálculo do total considerando múltiplos itens e o frete único
                valor_itens = (itens_completos['valor_final'].astype(float) * itens_completos['qtd'].astype(float)).sum()
                venda_total = valor_itens + float(dados_venda.iloc[0]['frete'])
                st.metric("Total da Proposta", f"R$ {venda_total:.2f}")

            st.write("**Itens que constarão no Word:**")
            st.dataframe(itens_completos[['sku_item', 'descricao', 'marca', 'qtd', 'valor_final']], use_container_width=True, hide_index=True)

        st.divider()

        # 3. Inputs Manuais Específicos da Proposta
        st.subheader("📝 Dados Adicionais da Proposta")
        with st.form("form_formalizacao"):
            f1, f2 = st.columns(2)
            with f1:
                n_pregao = st.text_input("Nº do Pregão / Processo", placeholder="Ex: 045/2024")
                validade = st.text_input("Validade da Proposta", value="60 (sessenta) dias")
            with f2:
                prazo = st.text_input("Prazo de Entrega", value="15 (quinze) dias úteis")
                contato_doc = st.text_input("Aos cuidados de:", value=cliente_info['nome_razao'])
            
            especificacoes = st.text_area("Informações Complementares / Especificações Técnicas")
            
            st.info("O sistema buscará o arquivo 'Proposta_Modelo.docx' na pasta do projeto.")
            botao_gerar = st.form_submit_button("🚀 GERAR ARQUIVO WORD", use_container_width=True)

        # 4. Lógica de Geração do Documento
        if botao_gerar:
            if not all([n_pregao, validade, prazo]):
                st.error("Por favor, preencha os campos de Pregão, Validade e Prazo.")
            else:
                try:
                    from num2words import num2words
                    from docx import Document
                    import io

                    # Carrega o modelo local
                    doc = Document("Proposta_Modelo.docx")
                    
                    # Valor por extenso para a Proposta
                    valor_extenso = num2words(venda_total, lang='pt_BR', to='currency').upper()

                    # Dicionário de Substituição (Tags [Tag] no seu Word)
                    subs = {
                        "[Razao_UASG]": str(cliente_info['nome_razao']),
                        "[N_pregao]": str(n_pregao),
                        "[Esp_solicitadas]": str(especificacoes),
                        "[Validade_Proposta]": str(validade),
                        "[Prazo_entrega]": str(prazo),
                        "[Endereco_Cliente]": f"{cliente_info['endereco']}, {cliente_info['numero']} - {cliente_info['bairro']}",
                        "[Contato_Cliente]": str(contato_doc),
                        "[Valor_Extenso]": valor_extenso,
                        "[Valor_Total]": f"R$ {venda_total:.2f}"
                    }

                    # Substitui no corpo do texto
                    for p in doc.paragraphs:
                        for tag, val in subs.items():
                            if tag in p.text:
                                p.text = p.text.replace(tag, val)

                    # Preenchimento automático da primeira tabela do Word
                    if doc.tables:
                        tabela = doc.tables[0]
                        for i, it in itens_completos.iterrows():
                            cells = tabela.add_row().cells
                            cells[0].text = str(i + 1)
                            cells[1].text = str(it['descricao'])
                            cells[2].text = str(it['marca'])
                            cells[3].text = str(it['sku_item']) # Usando SKU como PartNumber
                            cells[4].text = str(it['qtd'])
                            cells[5].text = f"R$ {float(it['valor_final']):.2f}"
                            cells[6].text = f"R$ {(float(it['valor_final']) * float(it['qtd'])):.2f}"

                    # Salva o arquivo em memória para download
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    st.success("✅ Proposta gerada com sucesso!")
                    st.download_button(
                        label="📥 Baixar Proposta (.docx)",
                        data=buffer,
                        file_name=f"Proposta_Pedido_{id_escolhido}_{cliente_info['nome_razao'][:15]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except FileNotFoundError:
                    st.error("Arquivo 'Proposta_Modelo.docx' não encontrado. Certifique-se de que ele está na mesma pasta do código.")
                except Exception as e:

                    st.error(f"Erro inesperado: {e}")







